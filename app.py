from flask import Flask, request, jsonify, render_template, session
import main
import json
import os
import re
from datetime import datetime
from collections import Counter
import speech_recognition as sr
import tempfile
import base64

app = Flask(__name__)
app.secret_key = '6598e01a2c5a5782d3241892e1fb99997a1dd6b043a4171e224bfe539184c7b0'

USERS_FILE = 'users.json'

def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, 'r') as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(USERS_FILE, 'w') as f:
        json.dump(users, f, indent=2)

# ─── NLP Analysis ─────────────────────────────────────────────────────────────

STOP_WORDS_ID = {
    'yang', 'dan', 'di', 'ke', 'dari', 'ini', 'itu', 'dengan', 'untuk',
    'adalah', 'pada', 'dalam', 'tidak', 'ada', 'saya', 'anda', 'kamu',
    'akan', 'sudah', 'atau', 'juga', 'bisa', 'lebih', 'telah', 'kami',
    'kita', 'mereka', 'apa', 'bagaimana', 'apakah', 'dapat', 'seperti',
    'karena', 'oleh', 'jika', 'maka', 'namun', 'saat', 'setiap', 'serta',
    'sehingga', 'sebuah', 'nya', 'pun', 'lah', 'kah', 'mu', 'ku', 'si',
    'para', 'pak', 'bu', 'bapak', 'ibu', 'dr', 'mr', 'mrs', 'hal', 'cara',
    'bagi', 'agar', 'mau', 'harus', 'boleh', 'punya', 'satu', 'dua', 'tiga',
    'dimana', 'siapa', 'kapan', 'mengapa', 'begitu', 'sampai', 'bahwa',
    'antara', 'setelah', 'sebelum', 'ketika', 'lalu', 'kemudian', 'masih',
    'itu', 'sangat', 'tapi', 'tetapi', 'kalau', 'cuma', 'hanya', 'paling'
}

STOP_WORDS_EN = {
    'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
    'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
    'should', 'may', 'might', 'shall', 'can', 'to', 'of', 'in', 'for',
    'on', 'with', 'at', 'by', 'from', 'as', 'into', 'through', 'during',
    'before', 'after', 'above', 'below', 'between', 'out', 'off', 'over',
    'under', 'again', 'then', 'once', 'i', 'me', 'my', 'we', 'our', 'you',
    'your', 'he', 'she', 'it', 'they', 'what', 'which', 'who', 'this',
    'that', 'these', 'those', 'and', 'but', 'or', 'nor', 'so', 'not',
    'if', 'about', 'up', 'just', 'very', 'also', 'all', 'any', 'both',
    'each', 'few', 'more', 'most', 'no', 'only', 'same', 'than', 'too',
    'when', 'where', 'why', 'how', 'some', 'such', 'there', 'here', 'its'
}

POS_WORDS_ID = {
    'bagus', 'baik', 'senang', 'suka', 'hebat', 'mantap', 'keren', 'puas',
    'benar', 'tepat', 'setuju', 'sukses', 'berhasil', 'mudah', 'sempurna',
    'cantik', 'indah', 'menarik', 'luar', 'biasa', 'terima', 'kasih',
    'luar biasa', 'oke', 'ya', 'tentu', 'jelas', 'bermanfaat', 'efektif',
    'cepat', 'akurat', 'lengkap', 'membantu', 'menyenangkan', 'positif'
}

NEG_WORDS_ID = {
    'buruk', 'jelek', 'salah', 'error', 'gagal', 'masalah', 'sulit',
    'susah', 'bingung', 'takut', 'kecewa', 'kurang', 'rusak', 'lambat',
    'malas', 'payah', 'lemah', 'tidak', 'belum', 'bahaya', 'berbahaya',
    'negatif', 'keliru', 'gagal', 'menyesal', 'khawatir', 'bosan',
    'marah', 'kecewa', 'sedih', 'pesimis', 'mustahil', 'percuma'
}

POS_WORDS_EN = {
    'good', 'great', 'excellent', 'happy', 'love', 'like', 'awesome',
    'perfect', 'wonderful', 'amazing', 'best', 'nice', 'cool', 'fantastic',
    'thank', 'thanks', 'yes', 'correct', 'right', 'easy', 'helpful',
    'brilliant', 'outstanding', 'superb', 'incredible', 'beautiful',
    'effective', 'efficient', 'accurate', 'useful', 'positive', 'clear'
}

NEG_WORDS_EN = {
    'bad', 'wrong', 'error', 'fail', 'problem', 'difficult', 'confused',
    'afraid', 'disappointed', 'sorry', 'broken', 'slow', 'worse', 'worst',
    'hate', 'dislike', 'terrible', 'awful', 'horrible', 'annoying',
    'frustrating', 'useless', 'impossible', 'failed', 'bug', 'issue',
    'sad', 'angry', 'boring', 'negative', 'incorrect', 'inaccurate'
}

ID_MARKERS = {
    'yang', 'dan', 'di', 'adalah', 'tidak', 'saya', 'anda', 'ini', 'itu',
    'dengan', 'untuk', 'apa', 'bagaimana', 'sudah', 'akan', 'atau', 'juga',
    'bisa', 'karena', 'jika', 'tetapi', 'namun', 'kami', 'kita', 'sangat'
}


def detect_language(words):
    id_count = sum(1 for w in words if w in ID_MARKERS)
    return 'Indonesian' if id_count >= 2 else 'English'


def analyze_nlp(text):
    if not text or not text.strip():
        return None

    text_clean = text.strip()
    words_raw = re.findall(r'\b\w+\b', text_clean.lower())
    sentences = [s.strip() for s in re.split(r'[.!?]+', text_clean) if s.strip()]

    # Remove stop words for meaningful analysis
    all_stops = STOP_WORDS_ID | STOP_WORDS_EN
    filtered_words = [w for w in words_raw if w not in all_stops and len(w) > 2]

    # Top words
    top_words = Counter(filtered_words).most_common(10)

    # Sentiment scoring
    pos_count = sum(1 for w in words_raw if w in POS_WORDS_ID | POS_WORDS_EN)
    neg_count = sum(1 for w in words_raw if w in NEG_WORDS_ID | NEG_WORDS_EN)

    total_sentiment_words = pos_count + neg_count
    if total_sentiment_words == 0:
        sentiment = 'neutral'
        sentiment_score = 0.0
        sentiment_pos = 0.0
        sentiment_neg = 0.0
    else:
        sentiment_pos = round(pos_count / total_sentiment_words, 3)
        sentiment_neg = round(neg_count / total_sentiment_words, 3)
        raw_score = (pos_count - neg_count) / total_sentiment_words
        sentiment_score = round(max(-1.0, min(1.0, raw_score)), 3)
        if sentiment_score > 0.1:
            sentiment = 'positive'
        elif sentiment_score < -0.1:
            sentiment = 'negative'
        else:
            sentiment = 'neutral'

    # Language
    language = detect_language(words_raw)

    # Stats
    word_count = len(words_raw)
    char_count = len(text_clean)
    sentence_count = max(len(sentences), 1)
    avg_word_length = round(sum(len(w) for w in words_raw) / max(word_count, 1), 2)
    unique_words = len(set(words_raw))
    lexical_diversity = round(unique_words / max(word_count, 1), 3)

    # Simple POS estimation
    # Nouns: words ending in typical noun suffixes
    noun_suffixes = ('tion', 'sion', 'ment', 'ness', 'ity', 'ism', 'ist',
                     'an', 'asi', 'kan', 'nya', 'an')
    verb_suffixes = ('ing', 'ed', 'ize', 'ise', 'ate',
                     'kan', 'mem', 'men', 'ber', 'ter', 'me')
    adj_suffixes = ('ful', 'less', 'ous', 'ive', 'able', 'ible',
                    'al', 'if', 'iah', 'is', 'ik')

    nouns_est = sum(1 for w in filtered_words if w.endswith(noun_suffixes))
    verbs_est = sum(1 for w in filtered_words if w.endswith(verb_suffixes))
    adjs_est = sum(1 for w in filtered_words if w.endswith(adj_suffixes))

    # Question detection
    is_question = text_clean.strip().endswith('?') or \
                  bool(re.match(r'^(apa|bagaimana|apakah|siapa|kapan|mengapa|kenapa|dimana|berapa|what|how|who|when|why|where|which|is|are|do|does|did|can|could|should|would)\b', text_clean.lower()))

    # Exclamation detection
    is_exclamation = text_clean.strip().endswith('!')

    return {
        'sentiment': sentiment,
        'sentiment_score': sentiment_score,
        'sentiment_pos': sentiment_pos,
        'sentiment_neg': sentiment_neg,
        'word_count': word_count,
        'char_count': char_count,
        'sentence_count': sentence_count,
        'avg_word_length': avg_word_length,
        'unique_words': unique_words,
        'lexical_diversity': lexical_diversity,
        'top_words': top_words,
        'language': language,
        'pos_keywords': pos_count,
        'neg_keywords': neg_count,
        'nouns_est': nouns_est,
        'verbs_est': verbs_est,
        'adjs_est': adjs_est,
        'is_question': is_question,
        'is_exclamation': is_exclamation,
        'filtered_word_count': len(filtered_words)
    }


# ─── Ekstraksi Teks dari File ─────────────────────────────────────────────────

def extract_text_from_pdf(file_path):
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except ImportError:
        pass
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        return f"[Gagal membaca PDF: {str(e)}]"


def extract_text_from_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[Gagal membaca DOCX: {str(e)}]"


def extract_text_from_excel(file_path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            result.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data):
                    result.append(" | ".join(row_data))
        return "\n".join(result)
    except Exception:
        pass
    try:
        import pandas as pd
        xl = pd.ExcelFile(file_path)
        result = []
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            result.append(f"=== Sheet: {sheet_name} ===")
            result.append(df.to_string(index=False))
        return "\n".join(result)
    except Exception as e:
        return f"[Gagal membaca Excel: {str(e)}]"


def extract_text_from_file(file_path, filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(file_path)
    elif ext in ('.xlsx', '.xls'):
        return extract_text_from_excel(file_path)
    elif ext in ('.txt', '.csv', '.md'):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        return None


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/nlp", methods=["GET"])
def nlp_page():
    return render_template("nlp.html")


@app.route("/api/nlp-analyze", methods=["POST"])
def nlp_analyze():
    data = request.get_json()
    text = data.get("text", "").strip()
    timestamp = data.get("timestamp", datetime.now().isoformat())
    session_id = data.get("session_id", "unknown")

    if not text:
        return jsonify({"error": "Teks tidak boleh kosong"}), 400

    result = analyze_nlp(text)
    if result is None:
        return jsonify({"error": "Gagal menganalisis teks"}), 500

    result['text_preview'] = text[:120] + ('...' if len(text) > 120 else '')
    result['timestamp'] = timestamp
    result['session_id'] = session_id

    return jsonify(result)


@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json()
    prompt = data.get("prompt", "")
    if not prompt:
        return jsonify({"error": "prompt tidak boleh kosong"}), 400
    result = main.generate_text(prompt)
    return jsonify({"prompt": prompt, "result": result})


@app.route("/generate-with-file", methods=["POST"])
def generate_with_file():
    prompt = request.form.get("prompt", "").strip()
    file = request.files.get("file")

    if not prompt:
        return jsonify({"error": "Prompt tidak boleh kosong"}), 400
    if not file or file.filename == "":
        return jsonify({"error": "Tidak ada file yang dikirim"}), 400

    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    allowed_exts = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.csv', '.md'}

    if ext not in allowed_exts:
        return jsonify({"error": f"Tipe file '{ext}' tidak didukung."}), 400

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        file_text = extract_text_from_file(tmp_path, filename)
        os.unlink(tmp_path)

        if file_text is None:
            return jsonify({"error": "Tipe file tidak didukung"}), 400
        if not file_text.strip():
            return jsonify({"error": "File tidak mengandung teks yang dapat dibaca"}), 400

        MAX_CHARS = 30000
        truncated = len(file_text) > MAX_CHARS
        if truncated:
            file_text = file_text[:MAX_CHARS] + "\n\n[... Teks dipotong karena terlalu panjang ...]"

        combined_prompt = (
            f"Berikut adalah isi file '{filename}':\n\n"
            f"---\n{file_text}\n---\n\n"
            f"Pertanyaan/Perintah pengguna: {prompt}"
        )

        result = main.generate_text(combined_prompt)
        return jsonify({"prompt": prompt, "filename": filename, "result": result, "truncated": truncated})

    except Exception as e:
        return jsonify({"error": f"Terjadi kesalahan saat memproses file: {str(e)}"}), 500


@app.route("/api/speech-to-text", methods=["POST"])
def speech_to_text():
    try:
        if 'audio' not in request.files:
            return jsonify({"error": "Tidak ada file audio"}), 400
        audio_file = request.files['audio']
        recognizer = sr.Recognizer()
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
            audio_file.save(temp_audio.name)
            with sr.AudioFile(temp_audio.name) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data, language='id-ID')
                os.unlink(temp_audio.name)
                return jsonify({"text": text})
    except sr.UnknownValueError:
        return jsonify({"error": "Tidak dapat memahami audio"}), 400
    except sr.RequestError as e:
        return jsonify({"error": f"Error dari service speech recognition: {e}"}), 500
    except Exception as e:
        return jsonify({"error": f"Terjadi kesalahan: {str(e)}"}), 500


@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json()
    name = data.get("name", "").strip()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    if not name or not email or not password:
        return jsonify({"error": "Semua field harus diisi"}), 400

    users = load_users()
    if email in users:
        return jsonify({"error": "Email sudah terdaftar"}), 400

    users[email] = {
        "id": f"user_{datetime.now().timestamp()}",
        "name": name, "email": email, "password": password,
        "created_at": datetime.now().isoformat()
    }
    save_users(users)
    return jsonify({"message": "Pendaftaran berhasil", "user": users[email]}), 201


@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    if not email or not password:
        return jsonify({"error": "Email dan password harus diisi"}), 400

    users = load_users()
    user = users.get(email)
    if not user or user["password"] != password:
        return jsonify({"error": "Email atau password salah"}), 401

    session["user_id"] = user["id"]
    session["user_email"] = user["email"]
    return jsonify({"message": "Login berhasil", "user": user})


@app.route("/api/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"message": "Logout berhasil"})


@app.route("/api/user", methods=["GET"])
def get_user():
    if "user_id" not in session:
        return jsonify({"error": "Not authenticated"}), 401
    users = load_users()
    user = users.get(session["user_email"])
    if not user:
        return jsonify({"error": "User not found"}), 404
    return jsonify({"user": user})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)